import io
import json
import os
import requests

from rest_framework import viewsets, status
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser

from django.http import StreamingHttpResponse

from document_generator.serializers import DocumentTemplateSerializer
from .models import DocumentTemplate

from docx import Document
from jinja2 import Environment

class GenerateDocumentViewSet(viewsets.ModelViewSet):
    serializer_class = DocumentTemplateSerializer
    queryset = DocumentTemplate.objects.all()

    def create(self, request, *args, **kwargs):
        parser_classes = [MultiPartParser]
        try:
            serializer = self.serializer_class(data = request.data)
            if serializer.is_valid():         
                #DATOS    
                template_text = serializer.data['template_text']
                data_text = serializer.data['data_text']
                format = serializer.data['format']
                #TEMPLATE POR TEXTO
                if template_text != "":
                    template = template_text
                # TEMPLATE POR DOCUMENTO
                else:
                    file_template = request.FILES['template_document'] 
                    # TEMPLATE POR DOCX
                    if file_template.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                        doc2 = Document(file_template)
                        template = "\n".join([p.text for p in doc2.paragraphs])
                    # TEMPLATE POR TXT
                    elif file_template.content_type == 'text/plain':
                        template = file_template.read().decode('utf-8')
                
                # DATA POR TEXTO
                if data_text != "":
                    data = eval(data_text)

                #DATA POR DOCUMENTO
                else:
                    file_data = request.FILES['data_document']
                    #DATA POR DOCX
                    if file_data.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                        doc3 = Document(file_data)
                        data = "\n".join([p.text for p in doc3.paragraphs])
                        data = eval(data)
                    #DATA POR JSON
                    elif file_data.content_type == 'application/json':
                        file_content = file_data.read() 
                        data = json.loads(file_content.decode('utf-8'))
                
                cont = 0
                response1 = []
                for a in data:
                    # RENDERIZACIÓN
                    env = Environment()
                    tmpl = env.from_string(template)
                    document_content = tmpl.render(a)
                    #SALIDA POR TEXTO PLANO
                    if format == 'text_plain':
                        response1.append({'document': document_content})
                        respuesta = Response(response1)
                    #SALIDA POR DOCX
                    elif format == 'docx':
                        doc = Document()
                        doc.add_paragraph(document_content)
                        carpeta_guardado = "../docx"

                        # Verificar si la carpeta existe, si no, crearla
                        if not os.path.exists(carpeta_guardado):
                            os.makedirs(carpeta_guardado)
                        url = os.path.join(carpeta_guardado, f"Test{cont}.docx")
                        with open(url, "wb") as f:
                            cont = cont + 1
                            doc.save(f)
                        
                        respuesta = Response({"messange":"Todo se guardo correctamente en tu carpeta DocumentGeneratorAPI"})
                      
                    else:
                        return Response({'error1': 'Invalid format'}, status=400)
                    
                return respuesta
            return Response({
                        'message': 'ha ocurrido un problema con serializer del producto',
                        'error':serializer.errors
                    }, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({'error2': str(e)}, status=500)
